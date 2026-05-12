from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Pt
from pathlib import Path


BASE_DIR = Path(r"D:\000 学习\05 信安\期末报告-微信社交迷思")
SOURCE_MARKDOWN = BASE_DIR / "报告正文_降噪版.md"
OUTPUT_PATH = BASE_DIR / "XXXXXX.姓名.期未报告.微信好友1000+，为什么真正聊得来的只有少数人：邓巴数视角下的数字社交迷思.docx"


def load_report_content() -> tuple[str, list[str]]:
    raw_text = SOURCE_MARKDOWN.read_text(encoding="utf-8").strip()
    lines = [line.rstrip() for line in raw_text.splitlines()]

    title = ""
    paragraphs: list[str] = []
    buffer: list[str] = []

    for line in lines:
        if line.startswith("# "):
            title = line[2:].strip()
            continue
        if not line.strip():
            if buffer:
                paragraphs.append("".join(buffer).strip())
                buffer = []
            continue
        buffer.append(line.strip())

    if buffer:
        paragraphs.append("".join(buffer).strip())

    if not title:
        raise ValueError("未在 Markdown 正文中找到一级标题。")
    if not paragraphs:
        raise ValueError("未在 Markdown 正文中找到正文段落。")

    return title, paragraphs


def set_default_font(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal_style.font.size = Pt(12)


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_title(document: Document) -> None:
    title_text, _ = load_report_content()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    title.paragraph_format.space_after = Pt(14)


def add_body(document: Document) -> None:
    _, body_paragraphs = load_report_content()
    for text in body_paragraphs:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = 1.5
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run(text)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def main() -> None:
    document = Document()
    configure_page(document)
    set_default_font(document)
    add_title(document)
    add_body(document)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
